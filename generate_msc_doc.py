#!/usr/bin/env python3
"""Generate Word document for MSc in Agentic AI and Innovative Application TPG program."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.expanduser("~/Desktop/MSc_Agentic_AI_培养方案.docx")
doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.8)
    s.right_margin = Cm(2.8)

# Colors
NAVY = RGBColor(0x0F, 0x17, 0x2A)
INDIGO = RGBColor(0x4F, 0x46, 0xE5)
DARK = RGBColor(0x1F, 0x2A, 0x37)
GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TEAL = RGBColor(0x0D, 0x94, 0x88)

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.font.color.rgb = DARK
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35


def shade(cell, color_hex):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:fill'), color_hex)
    sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)


def run_styled(p, text, size=10.5, bold=False, color=None, italic=False):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = '微软雅黑'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: r.bold = True
    if italic: r.italic = True
    if color: r.font.color.rgb = color
    return r


def heading(text, level=1):
    p = doc.add_paragraph()
    sizes = {1: 22, 2: 16, 3: 13, 4: 11.5}
    colors = {1: NAVY, 2: INDIGO, 3: DARK, 4: DARK}
    space_before = {1: 24, 2: 18, 3: 12, 4: 8}
    p.paragraph_format.space_before = Pt(space_before.get(level, 12))
    p.paragraph_format.space_after = Pt(8 if level <= 2 else 6)
    run_styled(p, text, size=sizes.get(level, 12), bold=True, color=colors.get(level, DARK))
    if level == 1:
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '12')
        bot.set(qn('w:color'), '4F46E5')
        bot.set(qn('w:space'), '6')
        pBdr.append(bot)
        pPr.append(pBdr)
    return p


def body(text, bold=False, color=None, indent=False, italic=False):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Cm(0.5)
    run_styled(p, text, bold=bold, color=color, italic=italic)
    return p


def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:color'), '4F46E5')
    left.set(qn('w:space'), '12')
    pBdr.append(left)
    pPr.append(pBdr)
    run_styled(p, text, size=11, bold=True, color=INDIGO)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    if bold_prefix:
        run_styled(p, '• ' + bold_prefix, bold=True)
        run_styled(p, text)
    else:
        run_styled(p, '• ' + text)
    return p


def table(headers, rows, col_widths=None, header_color='0F172A'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_styled(p, h, size=10, bold=True, color=WHITE)
        shade(c, header_color)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = ''
            run_styled(c.paragraphs[0], str(val), size=9.5)
            if ri % 2 == 1:
                shade(c, 'F9FAFB')
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def course_block(code, name_en, name_cn, credits, desc, topics=None, prereq=None):
    heading(f'{code}  {name_en}', 3)
    p = doc.add_paragraph()
    run_styled(p, name_cn, size=10.5, color=GRAY, italic=True)
    run_styled(p, f'    [{credits}学分]', size=10, bold=True, color=INDIGO)
    body(desc)
    if topics:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.5)
        run_styled(p2, '核心主题：', bold=True, color=DARK)
        run_styled(p2, topics, color=GRAY)
    if prereq:
        p3 = doc.add_paragraph()
        p3.paragraph_format.left_indent = Cm(0.5)
        run_styled(p3, '先修要求：', bold=True, color=DARK)
        run_styled(p3, prereq, color=GRAY)


# ══════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════

for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_styled(p, 'Master of Science in', size=18, color=GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_styled(p, 'Agentic AI and', size=36, bold=True, color=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_styled(p, 'Innovative Application', size=36, bold=True, color=INDIGO)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run_styled(p, '智能体及其创新应用', size=22, color=GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
run_styled(p, '理学硕士', size=20, color=GRAY)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_styled(p, 'Taught Postgraduate Programme', size=14, color=GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_styled(p, '教学型研究生培养方案', size=14, color=GRAY)

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_styled(p, '30学分 · 全日制1年 / 兼读制2年', size=12, color=GRAY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_styled(p, '2026–2027学年', size=12, color=GRAY)

doc.add_page_break()

# ══════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════

heading('目  录', 1)
toc_items = [
    '一、项目概述',
    '二、培养目标与预期学习成果',
    '三、入学要求',
    '四、学制与学分要求',
    '五、课程架构总览',
    '六、教学日历与学期安排',
    '七、核心课程详述',
    '八、选修课程详述',
    '九、毕业论文/顶点项目',
    '十、考核与毕业要求',
    '十一、师资与教学资源',
    '十二、职业发展方向',
]
for item in toc_items:
    body(item)

doc.add_page_break()

# ══════════════════════════════════════
# 1. 项目概述
# ══════════════════════════════════════

heading('一、项目概述', 1)

quote('培养掌握AI Agent核心技术、具备创新应用能力的高层次复合型人才')

body('理学硕士（智能体及其创新应用）项目 [MSc(AGIA)] 是一个以教学为主的授课型研究生项目，聚焦于人工智能领域最前沿的智能体（AI Agent）技术及其在各行业的创新应用。')

body('随着大语言模型（LLM）技术的突破性进展，AI Agent——能够自主感知环境、制定计划、使用工具并完成复杂任务的智能系统——正在成为人工智能的下一个核心范式。本项目区别于传统AI/机器学习硕士项目，专注于培养学生在Agent架构设计、多智能体协作、知识增强生成（RAG）、工具集成、以及AI驱动创新创业等方面的理论素养与实践能力。')

heading('项目特色', 2)

bullet('Agent-First 课程体系：', '以AI Agent为主线贯穿全部课程，而非传统的ML/DL→NLP/CV分支路径')
bullet('理论与实践并重：', '每门课程包含30-40%的实操项目，毕业前完成完整的Agent系统开发')
bullet('跨学科应用导向：', '选修课覆盖医疗、金融、教育、创意产业等垂直领域')
bullet('产业深度衔接：', '顶点项目可与企业合作，解决真实业务问题')
bullet('创新创业融合：', '设置AI创业方向选修课，培养技术+商业复合能力')

doc.add_page_break()

# ══════════════════════════════════════
# 2. 培养目标
# ══════════════════════════════════════

heading('二、培养目标与预期学习成果', 1)

heading('培养目标', 2)
body('本项目旨在培养能够独立设计、开发和部署AI Agent系统，并能将其创新应用于各行业场景的高层次专业人才。毕业生应具备：')
bullet('扎实的人工智能理论基础，特别是大语言模型、强化学习和知识表示')
bullet('熟练的Agent系统架构设计与工程实现能力')
bullet('多智能体系统的设计与协调能力')
bullet('将AI Agent技术应用于特定行业的创新实践能力')
bullet('AI伦理意识和负责任的AI开发理念')
bullet('独立研究和终身学习的能力')

heading('预期学习成果 (Programme Learning Outcomes)', 2)

table(
    ['编号', '学习成果', '对应课程'],
    [
        ['PLO1', '掌握AI与机器学习的核心理论，理解大语言模型的架构与训练方法', 'AGIA5001, AGIA5002'],
        ['PLO2', '能够设计和实现具备感知、推理、规划和行动能力的AI Agent系统', 'AGIA5003, AGIA5004'],
        ['PLO3', '理解多智能体系统的协作机制，能设计复杂的多Agent工作流', 'AGIA5004, 选修课'],
        ['PLO4', '掌握RAG、工具集成、记忆管理等Agent关键技术', 'AGIA5003, AGIA5101'],
        ['PLO5', '能够将AI Agent技术创新应用于至少一个垂直行业领域', '选修Group B, AGIA6001'],
        ['PLO6', '具备AI伦理意识，能够评估和管控AI系统的风险', 'AGIA5005'],
        ['PLO7', '具备独立完成AI Agent项目的全栈能力（设计→开发→部署→评估）', 'AGIA6001'],
    ],
    col_widths=[1.5, 8.5, 5.5],
)

doc.add_page_break()

# ══════════════════════════════════════
# 3. 入学要求
# ══════════════════════════════════════

heading('三、入学要求', 1)

heading('基本要求', 2)
bullet('持有认可大学的学士学位（任何学科），平均成绩不低于B或同等水平')
bullet('英语能力：TOEFL ≥ 80（iBT）或 IELTS ≥ 6.5，或同等英语资格')

heading('优先考虑', 2)
bullet('计算机科学、软件工程、数学、统计学、电子工程等相关学科背景')
bullet('具备Python编程经验（非硬性要求，Module 1将提供编程基础补充）')
bullet('有AI/数据科学相关项目或工作经验者优先')

heading('申请材料', 2)
bullet('本科成绩单及学位证书')
bullet('英语成绩证明')
bullet('个人陈述（Statement of Purpose, 500-800字）')
bullet('两封推荐信')
bullet('简历/CV')
bullet('作品集或项目链接（如有）')

heading('申请时间', 2)
table(
    ['轮次', '截止日期', '备注'],
    [
        ['第一轮', '2025年11月1日', '优先录取轮'],
        ['第二轮', '2026年1月15日', '主轮录取'],
        ['第三轮', '2026年3月15日', '补录轮（视学位余量）'],
    ],
    col_widths=[3, 4.5, 8],
)

doc.add_page_break()

# ══════════════════════════════════════
# 4. 学制与学分
# ══════════════════════════════════════

heading('四、学制与学分要求', 1)

table(
    ['项目', '全日制 (Full-time)', '兼读制 (Part-time)'],
    [
        ['标准学制', '1年', '2年'],
        ['最长修读期限', '2年', '4年'],
        ['每学期最低学分', '9学分', '6学分'],
        ['每学期最高学分', '15学分', '9学分'],
        ['总学分要求', '30学分', '30学分'],
    ],
    col_widths=[5, 5, 5.5],
)

heading('学分构成', 2)

table(
    ['类别', '学分', '课程数', '说明'],
    [
        ['核心课程 (Core)', '15', '5门', '必修，涵盖AI基础到Agent高级主题'],
        ['选修课程 Group A\n（Agent技术方向）', '6', '2门（最少）', '深化Agent技术栈'],
        ['选修课程 Group B\n（创新应用方向）', '3', '1门（最少）', '垂直行业应用'],
        ['顶点项目 (Capstone)', '6', '1门', '独立/小组项目，可与企业合作'],
        ['合计', '30', '9门 + 顶点项目', ''],
    ],
    col_widths=[4.5, 1.5, 3, 6.5],
)

doc.add_page_break()

# ══════════════════════════════════════
# 5. 课程架构总览
# ══════════════════════════════════════

heading('五、课程架构总览', 1)

table(
    ['课程编号', '课程名称（英文）', '课程名称（中文）', '学分', '类别'],
    [
        ['AGIA 5001', 'Foundations of AI and Machine Learning', 'AI与机器学习基础', '3', '核心'],
        ['AGIA 5002', 'Large Language Models', '大语言模型：架构与实践', '3', '核心'],
        ['AGIA 5003', 'AI Agent Design and Development', 'AI Agent设计与开发', '3', '核心'],
        ['AGIA 5004', 'Multi-Agent Systems', '多智能体系统与协作AI', '3', '核心'],
        ['AGIA 5005', 'AI Ethics, Safety and Governance', 'AI伦理、安全与治理', '3', '核心'],
        ['AGIA 5101', 'Knowledge Engineering and RAG', '知识工程与RAG系统', '3', '选修A'],
        ['AGIA 5102', 'NLP for Intelligent Agents', '面向智能体的自然语言处理', '3', '选修A'],
        ['AGIA 5103', 'Reinforcement Learning for Agents', '智能体强化学习', '3', '选修A'],
        ['AGIA 5104', 'Human-Agent Interaction', '人机智能体交互设计', '3', '选修A'],
        ['AGIA 5201', 'AI Agents for Business Innovation', 'AI Agent与商业创新创业', '3', '选修B'],
        ['AGIA 5202', 'AI Agents in Healthcare', 'AI Agent在医疗健康中的应用', '3', '选修B'],
        ['AGIA 5203', 'AI Agents in FinTech', 'AI Agent在金融科技中的应用', '3', '选修B'],
        ['AGIA 5204', 'AI Agents in Education', 'AI Agent在教育科技中的应用', '3', '选修B'],
        ['AGIA 5205', 'AI Agents for Creative Industries', 'AI Agent在创意产业中的应用', '3', '选修B'],
        ['AGIA 6000', 'Special Topics in Agentic AI', '智能体专题', '3', '选修'],
        ['AGIA 6001', 'Capstone Project', '顶点项目', '6', '必修'],
    ],
    col_widths=[2.3, 5.2, 4.5, 1, 1.5],
)

doc.add_page_break()

# ══════════════════════════════════════
# 6. 教学日历
# ══════════════════════════════════════

heading('六、教学日历与学期安排', 1)

heading('全日制学生建议修读计划', 2)

heading('秋季学期 Semester 1（9月 – 12月）', 3)
table(
    ['课程编号', '课程名称', '学分', '每周课时', '授课形式'],
    [
        ['AGIA 5001', 'AI与机器学习基础', '3', '3小时', '讲授2h + 实验1h'],
        ['AGIA 5002', '大语言模型：架构与实践', '3', '3小时', '讲授2h + 实验1h'],
        ['AGIA 5003', 'AI Agent设计与开发', '3', '3小时', '讲授1.5h + 实验1.5h'],
        ['选修A/B', '选修课1', '3', '3小时', '视课程而定'],
    ],
    col_widths=[2.3, 5, 1.2, 2, 5],
)
body('小计：12学分', bold=True, color=INDIGO)

heading('春季学期 Semester 2（1月 – 5月）', 3)
table(
    ['课程编号', '课程名称', '学分', '每周课时', '授课形式'],
    [
        ['AGIA 5004', '多智能体系统与协作AI', '3', '3小时', '讲授1.5h + 实验1.5h'],
        ['AGIA 5005', 'AI伦理、安全与治理', '3', '3小时', '讲授2h + 研讨1h'],
        ['选修A/B', '选修课2', '3', '3小时', '视课程而定'],
        ['选修A/B', '选修课3', '3', '3小时', '视课程而定'],
    ],
    col_widths=[2.3, 5, 1.2, 2, 5],
)
body('小计：12学分', bold=True, color=INDIGO)

heading('夏季学期 Summer Term（6月 – 8月）', 3)
table(
    ['课程编号', '课程名称', '学分', '周期', '形式'],
    [
        ['AGIA 6001', '顶点项目', '6', '10-12周', '独立研究 + 导师指导 + 答辩'],
    ],
    col_widths=[2.3, 5, 1.2, 2, 5],
)
body('小计：6学分', bold=True, color=INDIGO)

heading('关键时间节点', 2)
table(
    ['时间', '事项'],
    [
        ['9月第1周', '新生入学注册、迎新活动、选课'],
        ['9月第2周', '秋季学期正式开课'],
        ['10月中旬', '秋季学期期中评估'],
        ['12月中旬', '秋季学期考试周'],
        ['12月下旬 – 1月初', '冬季假期'],
        ['1月第2周', '春季学期正式开课'],
        ['2月底', '顶点项目选题截止'],
        ['3月中旬', '春季学期期中评估'],
        ['5月中旬', '春季学期考试周'],
        ['6月初', '夏季学期（顶点项目）开始'],
        ['7月中旬', '顶点项目中期检查'],
        ['8月中旬', '顶点项目答辩'],
        ['8月底', '成绩公布'],
        ['11月', '毕业典礼'],
    ],
    col_widths=[3.5, 12],
)

heading('兼读制学生建议修读计划', 2)

table(
    ['学期', '课程', '学分'],
    [
        ['Year 1 秋季', 'AGIA 5001 + AGIA 5002', '6'],
        ['Year 1 春季', 'AGIA 5003 + AGIA 5004', '6'],
        ['Year 1 夏季', '选修课1', '3'],
        ['Year 2 秋季', 'AGIA 5005 + 选修课2', '6'],
        ['Year 2 春季', '选修课3 + AGIA 6001（开始）', '3 + 6'],
        ['Year 2 夏季', 'AGIA 6001（答辩）', '—'],
    ],
    col_widths=[3.5, 7.5, 4.5],
)

doc.add_page_break()

# ══════════════════════════════════════
# 7. 核心课程详述
# ══════════════════════════════════════

heading('七、核心课程详述（15学分）', 1)

course_block(
    'AGIA 5001', 'Foundations of AI and Machine Learning', 'AI与机器学习基础', '3',
    '本课程为学生建立坚实的人工智能与机器学习理论基础。从AI的历史演进出发，系统介绍搜索算法、知识表示、概率推理、监督学习、无监督学习和深度学习等核心概念。课程注重直觉理解与数学原理的平衡，确保不同背景的学生都能建立完整的AI知识框架。',
    'AI历史与范式演进；搜索与优化算法；概率图模型与贝叶斯推理；监督学习（回归、分类、SVM、决策树）；无监督学习（聚类、降维）；深度学习基础（CNN、RNN）；模型评估与选择；Python机器学习工具栈（scikit-learn、PyTorch）',
    '基本数学素养（线性代数、微积分、概率论），建议具备Python编程经验'
)

course_block(
    'AGIA 5002', 'Large Language Models: Architecture and Practice', '大语言模型：架构与实践', '3',
    '本课程深入探讨大语言模型的核心技术，从Transformer架构到预训练、微调和对齐技术。学生将掌握Prompt Engineering的系统方法论，理解GPT、Claude、Llama等主流模型的设计哲学与能力边界，并通过API实践构建基于LLM的应用。',
    'Transformer架构深度解析；注意力机制与位置编码；预训练目标与训练流程；指令微调（SFT）与RLHF对齐；Prompt Engineering系统方法论（角色设定、思维链、Few-shot）；LLM API编程（OpenAI/Anthropic SDK）；Token经济与成本优化；开源vs闭源模型生态（GPT-4o/Claude/Llama/Qwen）；LLM的能力边界：幻觉、推理局限、知识截止',
    'AGIA 5001（可同步修读）'
)

course_block(
    'AGIA 5003', 'AI Agent Design and Development', 'AI Agent设计与开发', '3',
    '本课程是项目的核心旗舰课程，系统讲授AI Agent的设计原理、架构模式和工程实践。从"什么是Agent"的基本概念出发，逐步深入到ReAct、Plan-and-Execute等架构模式，涵盖工具调用（Function Calling）、记忆管理、RAG集成、状态管理等关键技术，并使用LangChain/LangGraph等主流框架进行实战开发。',
    'Agent定义与核心循环（Observe→Think→Act→Reflect）；Agent架构模式：ReAct、Plan-and-Execute、Tree-of-Thought；Function Calling与工具集成（MCP协议）；Agent记忆系统：短期/长期/工作记忆；RAG（检索增强生成）：向量数据库、分块策略、检索优化；LangChain/LangGraph框架实战；Agent评估方法与质量指标；Agent产品化：Streamlit/Gradio前端、API部署',
    'AGIA 5002'
)

course_block(
    'AGIA 5004', 'Multi-Agent Systems and Collaborative AI', '多智能体系统与协作AI', '3',
    '本课程探讨多个AI Agent如何协作完成单个Agent无法独立完成的复杂任务。从分布式AI的理论基础出发，介绍Agent间通信协议、任务分配策略、冲突解决机制，并通过CrewAI、AutoGen等框架实现多Agent协作系统。',
    '多Agent系统理论基础：分布式计算、博弈论基础；Agent间通信：消息传递、共享状态、黑板模型；协作模式：层级式、辩论式、流水线式、自组织式；任务分解与分配算法；CrewAI/AutoGen/MetaGPT框架实战；Agent Swarm与涌现行为；企业级多Agent架构设计；案例：AI软件开发团队、AI研究团队、AI客服中心',
    'AGIA 5003'
)

course_block(
    'AGIA 5005', 'AI Ethics, Safety and Governance', 'AI伦理、安全与治理', '3',
    '本课程从技术、哲学和政策三个维度审视AI系统的伦理问题、安全风险和治理框架。特别关注AI Agent作为自主决策系统所带来的独特挑战，如Agent对齐、Agent安全边界、自主权限管理等。',
    'AI伦理基础：公平性、透明性、可解释性、问责制；AI偏见：数据偏见、算法偏见的检测与缓解；Agent安全：Prompt注入防御、工具权限管控、自主边界设定；AI隐私：数据保护、联邦学习、差分隐私；AI治理框架：欧盟AI法案、中国AI治理、行业自律；AI对齐问题：目标对齐、价值对齐、RLHF的局限；负责任AI开发实践：Red Teaming、安全评估、部署监控',
    '无'
)

doc.add_page_break()

# ══════════════════════════════════════
# 8. 选修课程
# ══════════════════════════════════════

heading('八、选修课程详述', 1)

heading('Group A — Agent技术方向（至少6学分）', 2)

course_block(
    'AGIA 5101', 'Knowledge Engineering and RAG Systems', '知识工程与RAG系统', '3',
    '深入探讨如何为AI Agent构建和管理知识系统。从知识表示理论出发，系统讲授RAG（检索增强生成）的完整技术栈，包括文档处理、向量嵌入、检索策略、生成优化，以及知识图谱与Agent的融合应用。',
    '知识表示与本体论；文档处理管道（加载、分块、清洗）；嵌入模型与向量数据库（ChromaDB、Pinecone、Weaviate）；检索策略：稠密检索、稀疏检索、混合检索、重排序；高级RAG：多跳推理、自适应检索、GraphRAG；知识图谱构建与查询；Agent知识库管理与更新策略',
    'AGIA 5003'
)

course_block(
    'AGIA 5102', 'NLP for Intelligent Agents', '面向智能体的自然语言处理', '3',
    '从Agent的视角重新审视NLP技术，重点关注Agent在理解、生成和对话管理方面的需求。涵盖对话系统设计、意图识别、情感分析、多语言处理等Agent开发中的核心NLP能力。',
    '文本表示与语义理解；对话系统架构：检索式vs生成式；意图识别与槽填充；情感分析与观点挖掘；文本摘要与信息抽取；多语言与跨语言处理；Agent对话管理与上下文维护',
    'AGIA 5002'
)

course_block(
    'AGIA 5103', 'Reinforcement Learning for Agents', '智能体强化学习', '3',
    '探讨强化学习在AI Agent决策中的应用。从马尔可夫决策过程到深度强化学习，从单Agent到多Agent强化学习，培养学生设计能从环境反馈中持续学习和优化的智能Agent。',
    '马尔可夫决策过程（MDP）；动态规划与蒙特卡洛方法；Q-Learning与SARSA；深度Q网络（DQN）；策略梯度方法（REINFORCE、PPO、A3C）；多Agent强化学习（MARL）；从人类反馈中学习（RLHF）；实战：训练游戏AI、机器人导航Agent',
    'AGIA 5001'
)

course_block(
    'AGIA 5104', 'Human-Agent Interaction Design', '人机智能体交互设计', '3',
    '研究人与AI Agent之间的交互设计原则和方法。从认知科学和HCI理论出发，探讨如何设计直观、高效、可信的Agent交互体验，包括对话界面、多模态交互、Agent拟人化、用户信任建立等。',
    '人机交互（HCI）理论基础；对话交互设计原则；多模态交互：语音、视觉、手势；Agent拟人化与用户期望管理；可解释AI：如何让Agent"解释自己"；用户信任与Agent透明度；Agent UX评估方法；无障碍设计与包容性AI',
    '无'
)

heading('Group B — 创新应用方向（至少3学分）', 2)

course_block(
    'AGIA 5201', 'AI Agents for Business Innovation', 'AI Agent与商业创新创业', '3',
    '探讨AI Agent如何重塑商业模式和创业生态。从OPC（One Person Company）理念到AI-native企业架构，涵盖AI驱动的产品设计、市场营销、运营自动化和商业模式创新。学生将在课程中完成一份完整的AI创业商业计划。',
    'OPC模式与AI赋能个体创业；AI-native商业模式设计；AI Agent运营自动化栈（客服、内容、数据分析）；No-code/Low-code Agent平台（Dify、Coze）；工作流自动化：n8n/Make + Agent联动；AI产品的定价、获客与增长；案例研究与商业计划撰写',
    '无'
)

course_block(
    'AGIA 5202', 'AI Agents in Healthcare', 'AI Agent在医疗健康中的应用', '3',
    '探讨AI Agent在医疗健康领域的前沿应用，包括辅助诊断Agent、药物研发Agent、患者管理Agent、医学知识问答系统等，并深入讨论医疗AI的合规性、安全性和伦理考量。',
    '医学影像分析Agent；临床决策支持系统；药物发现与分子设计Agent；电子病历（EHR）智能处理；患者对话Agent与健康管理；医疗AI合规：HIPAA、FDA审批流程；医学知识图谱与循证医学',
    'AGIA 5003'
)

course_block(
    'AGIA 5203', 'AI Agents in FinTech', 'AI Agent在金融科技中的应用', '3',
    '探讨AI Agent在金融行业的创新应用，包括智能投顾Agent、风险评估Agent、反欺诈Agent、合规Agent等，以及金融大模型的训练与部署。',
    '金融数据处理与时序分析；智能投顾Agent设计；风险评估与信用评分Agent；反欺诈与异常检测；合规自动化Agent；算法交易与量化策略；金融大模型（FinGPT）；监管科技（RegTech）',
    'AGIA 5003'
)

course_block(
    'AGIA 5204', 'AI Agents in Education', 'AI Agent在教育科技中的应用', '3',
    '探讨AI Agent如何变革教育模式，包括个性化学习Agent、智能辅导系统、自适应评估、教学内容生成等，以及教育AI的伦理边界。',
    '学习科学与认知理论；个性化学习Agent设计；智能辅导系统（ITS）架构；自适应测评与学习路径规划；教育内容自动生成；MOOC平台AI集成；教育AI伦理：学术诚信、数据隐私、公平性',
    'AGIA 5003'
)

course_block(
    'AGIA 5205', 'AI Agents for Creative Industries', 'AI Agent在创意产业中的应用', '3',
    '探讨AI Agent在内容创作、设计、音乐、游戏等创意产业中的应用。涵盖多模态生成（文本、图像、音频、视频）、AI辅助创意工作流、版权与知识产权等议题。',
    '多模态生成模型（DALL-E、Midjourney、Sora）；AI辅助写作与内容创作Agent；AI音乐生成与编曲；游戏AI：NPC Agent、程序化内容生成；AI辅助设计工作流；创意产业的版权与IP问题；人机协作创作模式',
    '无'
)

course_block(
    'AGIA 6000', 'Special Topics in Agentic AI', '智能体专题', '3',
    '本课程涵盖AI Agent领域的前沿专题，内容每年更新以反映最新技术进展。可能的主题包括：具身智能体（Embodied Agent）、世界模型、Agent安全对齐、自主Agent系统等。本课程可重复修读（不同主题）。',
    '视当年主题而定',
    '视主题而定'
)

doc.add_page_break()

# ══════════════════════════════════════
# 9. 顶点项目
# ══════════════════════════════════════

heading('九、毕业论文 / 顶点项目', 1)

course_block(
    'AGIA 6001', 'Capstone Project', '顶点项目', '6',
    '顶点项目是本项目的核心实践环节，要求学生综合运用所学知识，独立或以小组（2-3人）形式完成一个完整的AI Agent系统。项目可以是学术研究型、产品开发型或企业合作型。',
    None, None
)

heading('项目类型', 3)
table(
    ['类型', '说明', '产出'],
    [
        ['学术研究型', '在导师指导下针对AI Agent领域的一个研究问题进行探索', '研究论文（8000-12000字）+ 原型系统'],
        ['产品开发型', '设计并开发一个完整的AI Agent应用产品', '可部署的产品 + 技术文档 + 用户测试报告'],
        ['企业合作型', '与合作企业共同解决一个真实业务问题', '解决方案 + 系统实现 + 企业评估报告'],
    ],
    col_widths=[3, 6, 6.5],
)

heading('时间节点', 3)
table(
    ['阶段', '时间', '任务', '产出'],
    [
        ['选题', '春季学期第6周前', '确定导师、选题、研究/开发计划', '项目提案（2000字）'],
        ['开题', '春季学期第10周', '开题报告答辩', '开题报告 + PPT'],
        ['开发', '夏季学期第1-8周', '核心开发/研究工作', '周进度报告'],
        ['中期检查', '夏季学期第5周', '中期进展汇报', '中期报告'],
        ['答辩', '夏季学期第10-11周', '最终答辩', '论文/报告 + 系统演示 + 答辩PPT'],
    ],
    col_widths=[2, 4, 4.5, 5],
)

heading('评分标准', 3)
table(
    ['维度', '权重', '说明'],
    [
        ['技术深度', '30%', 'Agent系统架构设计的合理性、技术实现的质量和创新性'],
        ['实践价值', '25%', '项目解决真实问题的能力、产品完成度、可部署性'],
        ['学术规范', '20%', '文献综述的全面性、方法论的严谨性、写作质量'],
        ['演示答辩', '15%', '口头表达能力、回答问题的深度、Demo展示效果'],
        ['过程管理', '10%', '进度管理、导师评价、周报完成情况'],
    ],
    col_widths=[3, 1.5, 11],
)

doc.add_page_break()

# ══════════════════════════════════════
# 10. 考核与毕业
# ══════════════════════════════════════

heading('十、考核与毕业要求', 1)

heading('课程考核方式', 2)
body('各课程采用多元化评估方式，典型构成如下：')
table(
    ['评估方式', '权重', '说明'],
    [
        ['课堂参与', '10%', '出勤、讨论、课堂练习'],
        ['作业与实验', '30%', '每2-3周一次的编程作业或实验报告'],
        ['期中项目', '20%', '中等规模的Agent开发项目'],
        ['期末考核', '40%', '期末考试（闭卷/开卷）或期末项目'],
    ],
    col_widths=[3, 1.5, 11],
)
body('注：具体权重由各课程教师确定，以课程大纲为准。', italic=True, color=GRAY)

heading('成绩评定', 2)
table(
    ['等级', '绩点', '百分制参考', '说明'],
    [
        ['A+', '4.3', '95-100', '优异'],
        ['A', '4.0', '90-94', '优秀'],
        ['A-', '3.7', '85-89', '良好偏上'],
        ['B+', '3.3', '80-84', '良好'],
        ['B', '3.0', '75-79', '中等偏上'],
        ['B-', '2.7', '70-74', '中等'],
        ['C+', '2.3', '65-69', '及格偏上'],
        ['C', '2.0', '60-64', '及格'],
        ['F', '0.0', '<60', '不及格'],
    ],
    col_widths=[2, 2, 3.5, 8],
)

heading('毕业要求', 2)
bullet('完成全部30学分课程')
bullet('累计GPA不低于2.85（B平均）')
bullet('顶点项目成绩不低于C+')
bullet('无学术不端记录')
bullet('在规定学制内完成所有要求')

doc.add_page_break()

# ══════════════════════════════════════
# 11. 师资
# ══════════════════════════════════════

heading('十一、师资与教学资源', 1)

heading('师资构成', 2)
bullet('全职教授/副教授：', '负责核心课程教学与研究指导')
bullet('业界导师：', '来自头部科技企业的AI专家，参与应用课程教学与顶点项目指导')
bullet('访问学者：', '国际知名AI研究机构的学者，开设Special Topics课程')

heading('教学资源', 2)
bullet('GPU计算集群：', '提供NVIDIA A100/H100计算资源，支持大模型微调和训练')
bullet('API资源：', '提供OpenAI、Anthropic、Google等主流LLM API学术额度')
bullet('软件许可：', 'GitHub Copilot、JetBrains IDE等开发工具的教育许可')
bullet('合作企业网络：', '与科技公司建立合作，提供实习和项目合作机会')
bullet('学术数据库：', 'IEEE、ACM、Springer等学术论文全文访问')

doc.add_page_break()

# ══════════════════════════════════════
# 12. 职业方向
# ══════════════════════════════════════

heading('十二、职业发展方向', 1)

body('本项目毕业生具备独特的AI Agent专业技能，可在以下方向发展：')

table(
    ['方向', '典型职位', '目标企业/机构'],
    [
        ['AI Agent工程师', 'Agent Developer, LLM Engineer, AI Platform Engineer', '科技公司、AI创业公司'],
        ['AI产品经理', 'AI Product Manager, Technical PM', '互联网公司、SaaS企业'],
        ['AI解决方案架构师', 'AI Solution Architect, Technical Consultant', '咨询公司、系统集成商'],
        ['AI创业者', 'Founder / CTO', '自主创业、AI孵化器'],
        ['AI研究员', 'Research Scientist, ML Engineer', '研究院、大厂AI Lab'],
        ['行业AI专家', '医疗AI / 金融AI / 教育AI专家', '垂直行业领军企业'],
        ['AI安全与治理', 'AI Safety Researcher, AI Policy Analyst', '政府机构、智库、大型企业'],
        ['深造', '博士研究生（AI/CS方向）', '全球顶尖大学'],
    ],
    col_widths=[3.5, 5.5, 6.5],
)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
# Decorative line
pPr = p._element.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
top = OxmlElement('w:top')
top.set(qn('w:val'), 'single')
top.set(qn('w:sz'), '6')
top.set(qn('w:color'), '4F46E5')
top.set(qn('w:space'), '12')
pBdr.append(top)
pPr.append(pBdr)
run_styled(p, 'MSc in Agentic AI and Innovative Application', size=11, color=GRAY, italic=True)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_styled(p2, '智能体及其创新应用 · 理学硕士 · 2026-2027', size=11, color=GRAY, italic=True)

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
