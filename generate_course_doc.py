#!/usr/bin/env python3
"""Generate a polished Word document for the AI Agent & OPC training program."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.expanduser("~/Desktop/AI_Agent_OPC_培训课程方案.docx")

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ── Style helpers ──
NAVY = RGBColor(0x1A, 0x23, 0x5B)
INDIGO = RGBColor(0x4F, 0x46, 0xE5)
DARK = RGBColor(0x1F, 0x2A, 0x37)
GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF3, 0xF4, 0xF6)
INDIGO_BG = RGBColor(0xEE, 0xF2, 0xFF)
GREEN_BG = RGBColor(0xEC, 0xFD, 0xF5)
AMBER_BG = RGBColor(0xFF, 0xFB, 0xEB)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
font.color.rgb = DARK
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.3


def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '4F46E5'))
        element.set(qn('w:space'), '0')
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_styled_heading(text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = NAVY
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        # Add bottom border line
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:color'), '4F46E5')
        bottom.set(qn('w:space'), '6')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = INDIGO
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    elif level == 3:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = DARK
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p


def add_body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    return p


def add_quote_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Add left border
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:color'), '4F46E5')
    left.set(qn('w:space'), '12')
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = INDIGO
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.bold = True
    return p


def make_table(headers, rows, col_widths=None, header_color='1A235B'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F9FAFB')

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table


def make_schedule_table(rows):
    """Create a week schedule table: 时段 | 内容 | 形式"""
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['时段', '内容', '形式']
    widths = [Cm(2), Cm(11), Cm(2.5)]

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        cell.width = widths[i]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4F46E5')

    for r_idx, (period, content, form) in enumerate(rows):
        for c_idx, val in enumerate([period, content, form]):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            cell.width = widths[c_idx]
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            if c_idx == 0:
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F3F4F6')

    doc.add_paragraph()
    return table


# ════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AI Agent & OPC')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = NAVY
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('创新应用培训班')
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = INDIGO
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run('课  程  方  案')
run.font.size = Pt(20)
run.font.color.rgb = GRAY
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in ['144学时 · 12周 · 6模块', '面向本科生 · 零基础入门到项目实战', '2026年3月']:
    run = p.add_run(line + '\n')
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ════════════════════════════════════════════
# 课程总览
# ════════════════════════════════════════════

add_styled_heading('课程总览', 1)

make_table(
    ['项目', '内容'],
    [
        ['总学时', '144学时'],
        ['授课周期', '12周（每周末上课，周六 + 周日）'],
        ['每日安排', '上午3小时（理论授课）+ 下午3小时（实操练习）'],
        ['模块数', '6个模块，每模块2个周末（24学时）'],
        ['面向对象', '本科生（无需AI基础）'],
        ['最终产出', '每组完成一个AI Agent驱动的OPC创业项目路演'],
    ],
    col_widths=[3.5, 12],
)

# ════════════════════════════════════════════
# MODULE 1
# ════════════════════════════════════════════

add_styled_heading('Module 1：AI基础与大模型认知', 1)
add_body('第1-2周 · 24学时', bold=True, color=GRAY)
add_quote_block('目标：从零建立AI认知，能熟练使用主流大模型工具')

add_styled_heading('Week 1 周六 — AI是什么？从历史到现在', 3)
make_schedule_table([
    ('上午', 'AI发展史：从图灵测试→深度学习→ChatGPT时刻；ML/DL/LLM概念辨析（不讲数学，讲直觉）；大模型能做什么、不能做什么（幻觉、知识截止、推理局限）', '讲授+演示'),
    ('下午', '实操：注册并使用ChatGPT、Claude、文心一言、Kimi；完成5个真实任务（写邮件、翻译、代码生成、数据分析、创意brainstorm）；对比不同模型的表现差异', '动手实验'),
])

add_styled_heading('Week 1 周日 — 大模型是怎么工作的', 3)
make_schedule_table([
    ('上午', 'Token与分词原理（动手看tokenizer）；Transformer直觉讲解（注意力机制≈"找重点"）；预训练→微调→RLHF流程（概念级）；开源vs闭源模型生态（GPT-4o / Claude / Llama / Qwen）', '讲授+可视化'),
    ('下午', '实操：用Hugging Face Playground体验不同模型；调节temperature/top_p参数观察输出变化；用API发送第一个请求（Python/curl）；小组任务：每组选一个行业，列出5个AI能解决的真实痛点', '动手实验'),
])

add_styled_heading('Week 2 周六 — Prompt Engineering（上）', 3)
make_schedule_table([
    ('上午', 'Prompt Engineering核心原则：角色设定、任务拆解、格式控制、Few-shot示例；Prompt框架：CRISPE、CO-STAR；从"写个方案"到精准输出的prompt迭代过程', '讲授+案例'),
    ('下午', '实操：10个prompt挑战赛（从简单到复杂）；Chain-of-Thought（思维链）实战；让LLM输出结构化JSON；作业：用prompt完成一个完整的市场分析报告', '竞赛+练习'),
])

add_styled_heading('Week 2 周日 — Prompt Engineering（下）+ API编程入门', 3)
make_schedule_table([
    ('上午', '高级prompt技巧：Multi-turn对话管理、System Prompt设计、Prompt Chaining（任务串联）；Token成本计算与优化；Prompt安全（注入攻击与防御）', '讲授+演示'),
    ('下午', '实操：Python基础速通（变量、函数、API调用——仅覆盖所需最少内容）；用OpenAI/Anthropic SDK写第一个AI应用；模块项目：构建一个"智能客服"命令行工具', '编程实操'),
])

add_body('📦 模块1产出：每位学员能独立调用LLM API，掌握prompt engineering，完成命令行AI工具', bold=True, color=INDIGO)

doc.add_page_break()

# ════════════════════════════════════════════
# MODULE 2
# ════════════════════════════════════════════

add_styled_heading('Module 2：AI Agent原理与基础开发', 1)
add_body('第3-4周 · 24学时', bold=True, color=GRAY)
add_quote_block('目标：理解Agent架构，能用框架构建基础AI Agent')

add_styled_heading('Week 3 周六 — 什么是AI Agent', 3)
make_schedule_table([
    ('上午', '从LLM到Agent的跨越：LLM只是"大脑"，Agent = 大脑+感知+行动+记忆；Agent核心循环：Observe→Think→Act→Reflect；经典架构：ReAct、Plan-and-Execute、Tool Use；案例拆解：Devin（AI程序员）、AutoGPT、ChatGPT Plugins', '讲授+案例'),
    ('下午', '实操：用纯Python（无框架）实现一个最简Agent循环；添加工具调用：让Agent能搜索网页、执行计算；理解Function Calling机制（OpenAI/Claude格式）；构建一个能查天气+做数学计算的Agent', '编程实操'),
])

add_styled_heading('Week 3 周日 — Agent开发框架', 3)
make_schedule_table([
    ('上午', '主流Agent框架对比：LangChain、LangGraph、CrewAI、AutoGen、Dify；框架选型原则（什么场景用什么框架）；LangChain核心概念：Chain、Tool、Memory、Agent', '讲授+演示'),
    ('下午', '实操：用LangChain构建一个完整Agent；添加网页搜索工具（Tavily/SerpAPI）；添加对话记忆（ConversationBufferMemory）；小组任务：构建一个"AI研究助手"——输入主题，自动搜索+总结+生成报告', '编程实操'),
])

add_styled_heading('Week 4 周六 — RAG：让Agent拥有专属知识', 3)
make_schedule_table([
    ('上午', '为什么需要RAG（大模型知识截止、幻觉问题）；RAG架构：文档加载→分块→嵌入→向量存储→检索→生成；Embedding原理（直觉讲解：文本→数字向量→语义相似度）；向量数据库：ChromaDB、FAISS、Pinecone', '讲授+架构图'),
    ('下午', '实操：构建完整RAG管道——加载PDF/网页→切分→存入ChromaDB→查询→AI回答；调参实验：chunk_size、overlap、top_k对效果的影响；评估RAG质量；动手：为一本教材构建"AI助教"', '编程实操'),
])

add_styled_heading('Week 4 周日 — Agent记忆与状态管理', 3)
make_schedule_table([
    ('上午', 'Agent记忆类型：短期（对话上下文）、长期（向量数据库）、工作记忆（scratchpad）；状态管理：有限状态机 vs 图结构（LangGraph）；Agent的错误处理与自我纠错', '讲授+案例'),
    ('下午', '实操：用LangGraph构建有状态的Agent工作流；实现Agent的"反思"能力——检查自己的输出并改进；模块项目：构建一个"智能文档问答系统"——上传文档→AI建立知识库→多轮问答', '编程实操'),
])

add_body('📦 模块2产出：每组完成一个RAG驱动的智能文档问答Agent', bold=True, color=INDIGO)

doc.add_page_break()

# ════════════════════════════════════════════
# MODULE 3
# ════════════════════════════════════════════

add_styled_heading('Module 3：Agent高级能力与多智能体系统', 1)
add_body('第5-6周 · 24学时', bold=True, color=GRAY)
add_quote_block('目标：掌握复杂Agent模式，理解多Agent协作')

add_styled_heading('Week 5 周六 — 工具集成与API编排', 3)
make_schedule_table([
    ('上午', 'Agent工具生态：代码执行、数据库查询、文件操作、邮件发送、日历管理；MCP（Model Context Protocol）：标准化的工具接入协议；API编排：让Agent串联多个外部服务（支付、地图、CRM）', '讲授+演示'),
    ('下午', '实操：为Agent添加5+工具（搜索、计算、代码执行、文件读写、数据库查询）；用MCP Server接入自定义工具；构建一个能"做事"的Agent：自动发邮件/生成报表/操作数据库', '编程实操'),
])

add_styled_heading('Week 5 周日 — 多智能体系统', 3)
make_schedule_table([
    ('上午', '为什么需要多Agent：单个Agent的能力上限；协作模式：层级式（Boss→Worker）、辩论式、流水线式；CrewAI/AutoGen框架实战；案例：AI软件开发团队（PM→Architect→Developer→Tester）', '讲授+案例'),
    ('下午', '实操：用CrewAI构建一个3-Agent团队——Researcher + Writer + Editor协作完成一篇文章；设计Agent间的通信协议和任务分配；小组任务：设计一个5-Agent协作系统解决一个真实问题', '编程实操'),
])

add_styled_heading('Week 6 周六 — Agent评估与优化', 3)
make_schedule_table([
    ('上午', '如何评估Agent质量：任务完成率、准确度、效率、成本；常见问题：Agent循环、工具调用失败、幻觉传播；优化策略：更好的System Prompt、工具描述优化、温度参数调节；成本控制：token用量监控、缓存策略、模型分级调用', '讲授+分析'),
    ('下午', '实操：为Module 2的项目添加评估指标；A/B测试不同prompt策略的效果；实现成本监控dashboard；动手：优化一个"差Agent"——诊断问题并修复', '实验+优化'),
])

add_styled_heading('Week 6 周日 — Agent产品化', 3)
make_schedule_table([
    ('上午', '从Demo到产品：Agent部署架构（前端+后端+LLM API）；前端界面：Streamlit/Gradio快速搭建；用户体验设计：流式输出、错误提示、加载状态；安全考量：输入过滤、输出审核、API密钥保护', '讲授+演示'),
    ('下午', '实操：用Streamlit将Agent包装成Web应用；添加用户认证和使用限制；部署到云端（Streamlit Cloud / Railway）；模块项目：将Module 2的Agent产品化——完整Web界面、可公开访问', '编程实操'),
])

add_body('📦 模块3产出：每组的Agent项目具备Web界面，可公开演示', bold=True, color=INDIGO)

doc.add_page_break()

# ════════════════════════════════════════════
# MODULE 4
# ════════════════════════════════════════════

add_styled_heading('Module 4：OPC模式与AI驱动创业', 1)
add_body('第7-8周 · 24学时', bold=True, color=GRAY)
add_quote_block('目标：理解OPC（One Person Company）模式，掌握AI如何赋能个体创业')

add_styled_heading('Week 7 周六 — OPC革命：一个人就是一家公司', 3)
make_schedule_table([
    ('上午', 'OPC概念：AI让一个人具备一个团队的能力；传统公司的职能拆解：市场、产品、开发、运营、客服、财务；AI Agent替代矩阵：哪些职能可以被Agent化？（80%可以）；案例研究：年收入百万的一人公司（实际案例3-5个）', '讲授+案例'),
    ('下午', '工作坊：每组选择一个创业方向，绘制"AI职能替代图"；头脑风暴：你的一人公司可以做什么？；用AI工具完成市场调研全流程（选题→搜索→分析→报告→PPT）；产出：每组完成一份AI辅助的商业机会分析', '工作坊'),
])

add_styled_heading('Week 7 周日 — AI时代的商业模式设计', 3)
make_schedule_table([
    ('上午', 'AI原生商业模式：SaaS→AI-aaS的转变；定价策略：按次/按月/按效果；获客渠道：内容营销（AI生成）、社群运营、产品驱动增长；成本结构：API调用费 vs 传统人力成本对比', '讲授+分析'),
    ('下午', '实操：用AI Agent完成一套完整商业计划书；构建自动化内容生产管道（选题→写作→配图→发布）；小组任务：设计你的OPC商业模式画布（Business Model Canvas）', '工作坊+实操'),
])

add_styled_heading('Week 8 周六 — AI运营工具栈', 3)
make_schedule_table([
    ('上午', 'OPC运营全栈：AI客服Agent、AI内容Agent、AI数据分析Agent、AI财务Agent；No-code/Low-code工具：Dify、Coze、FastGPT搭建Agent；工作流自动化：n8n/Make + AI Agent联动；个人品牌构建：AI辅助的全渠道内容策略', '讲授+演示'),
    ('下午', '实操：在Dify/Coze上搭建一个完整的业务Agent（不写代码）；用n8n搭建自动化工作流（新客户→AI回复→记录到表格→提醒）；动手：为你的OPC项目搭建一个AI客服系统', '动手实验'),
])

add_styled_heading('Week 8 周日 — 真实案例深度拆解 + 项目选题', 3)
make_schedule_table([
    ('上午', '5个成功OPC案例深度拆解：技术栈、商业模式、增长路径、收入结构；失败案例分析：为什么有些AI创业失败了？；行业机会地图：教育、医疗、法律、电商、内容、金融', '案例教学'),
    ('下午', '期末项目选题：每组确定一个AI Agent + OPC创业项目；撰写项目提案（问题→方案→技术→商业模式→团队分工）；导师点评与反馈；产出：确定最终项目方向，提交项目提案', '工作坊'),
])

add_body('📦 模块4产出：每组确定创业项目方向，完成商业计划书和项目提案', bold=True, color=INDIGO)

doc.add_page_break()

# ════════════════════════════════════════════
# MODULE 5
# ════════════════════════════════════════════

add_styled_heading('Module 5：综合项目开发', 1)
add_body('第9-10周 · 24学时', bold=True, color=GRAY)
add_quote_block('目标：动手构建一个完整的AI Agent驱动的OPC创业项目')

add_styled_heading('Week 9 周六 — 项目架构设计与核心开发', 3)
make_schedule_table([
    ('上午', '项目架构设计指导：前端+后端+Agent+数据库的整体架构；技术选型建议（根据各组项目）；MVP（最小可行产品）思维：先做核心功能', '指导+设计'),
    ('下午', '分组开发：核心Agent逻辑实现；导师巡回指导', '项目开发'),
])

add_styled_heading('Week 9 周日 — 核心功能完成', 3)
make_schedule_table([
    ('上午', '技术答疑专场：各组遇到的技术难题集中解决；代码审查与优化建议', '答疑+代码审查'),
    ('下午', '分组开发：完成MVP核心功能；阶段演示：各组展示当前进度（5分钟/组）', '项目开发+演示'),
])

add_styled_heading('Week 10 周六 — 产品打磨与集成', 3)
make_schedule_table([
    ('上午', '产品打磨技巧：UI/UX优化、错误处理、边界情况；数据与隐私：用户数据处理规范', '讲授+指导'),
    ('下午', '分组开发：前端界面完善、功能集成测试；用户体验测试（组间交叉测试）', '项目开发'),
])

add_styled_heading('Week 10 周日 — 部署上线与运营准备', 3)
make_schedule_table([
    ('上午', '部署实战：云服务器配置、域名绑定、HTTPS；运营准备：Landing Page制作、用户反馈收集机制', '实操'),
    ('下午', '完成部署上线；准备项目演示材料；产出：项目可在线访问', '项目开发'),
])

add_body('📦 模块5产出：每组项目MVP完成并部署上线', bold=True, color=INDIGO)

doc.add_page_break()

# ════════════════════════════════════════════
# MODULE 6
# ════════════════════════════════════════════

add_styled_heading('Module 6：路演与未来展望', 1)
add_body('第11-12周 · 24学时', bold=True, color=GRAY)
add_quote_block('目标：完善项目，完成路演，理解AI行业未来方向')

add_styled_heading('Week 11 周六 — 路演技巧与商业包装', 3)
make_schedule_table([
    ('上午', '路演技巧：如何在5分钟内讲清楚一个AI产品；商业故事：问题→方案→市场→竞品→团队→融资；Demo展示技巧：如何避免现场翻车', '讲授+观摩'),
    ('下午', '路演PPT制作（AI辅助）；第一轮模拟路演（组间互评）；根据反馈迭代改进', '练习+反馈'),
])

add_styled_heading('Week 11 周日 — 项目优化与预演', 3)
make_schedule_table([
    ('上午', '最后的产品优化窗口；用户测试数据收集和展示准备', '项目优化'),
    ('下午', '完整预演（每组15分钟：10分钟演示+5分钟答辩）；导师深度反馈；最终调整', '预演+反馈'),
])

add_styled_heading('Week 12 周六 — 正式路演', 3)
make_schedule_table([
    ('上午', '正式路演（每组15分钟）；评审团打分（技术可行性30%、商业价值30%、产品完成度20%、展示表现20%）', '正式路演'),
    ('下午', '评审点评与颁奖；优秀项目后续孵化计划介绍', '点评+颁奖'),
])

add_styled_heading('Week 12 周日 — AI前沿与职业发展', 3)
make_schedule_table([
    ('上午', 'AI前沿趋势：具身智能、世界模型、AGI进展；AI时代的职业规划：哪些能力不会被替代；终身学习路线图：从本课程出发的进阶路径', '讲座'),
    ('下午', '校友网络建立；项目后续合作对接；结业仪式；课程回顾与总结', '交流+结业'),
])

add_body('📦 模块6产出：完成路演，获得评审反馈，结业', bold=True, color=INDIGO)

doc.add_page_break()

# ════════════════════════════════════════════
# 能力成长路线
# ════════════════════════════════════════════

add_styled_heading('课程能力成长路线', 1)

make_table(
    ['阶段', '角色定位', '核心能力'],
    [
        ['Module 1', 'AI 用户', '会用ChatGPT、会写Prompt、能调API'],
        ['Module 2', 'AI 开发者', '能建Agent、能做RAG、能管理记忆与状态'],
        ['Module 3', 'AI 架构师', '能搭系统、能建多Agent协作、能部署上线'],
        ['Module 4', 'AI 创业者', '能设计商业模式、能算ROI、能获客运营'],
        ['Module 5-6', 'AI 产品人', '能交付完整产品、能路演展示、能持续迭代'],
    ],
    col_widths=[3, 3, 9.5],
)

# ════════════════════════════════════════════
# 考核方式
# ════════════════════════════════════════════

add_styled_heading('考核方式', 1)

make_table(
    ['项目', '占比', '说明'],
    [
        ['平时作业', '20%', '每模块一个小作业'],
        ['模块项目', '30%', 'Module 2 + Module 3 的技术项目'],
        ['期末路演', '40%', '完整的AI + OPC创业项目路演'],
        ['课堂参与', '10%', '出勤、讨论、互评'],
    ],
    col_widths=[3, 2, 10.5],
)

# ════════════════════════════════════════════
# Save
# ════════════════════════════════════════════

doc.save(OUTPUT_PATH)
print(f"Document saved to: {OUTPUT_PATH}")
