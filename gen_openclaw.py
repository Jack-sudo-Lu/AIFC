#!/usr/bin/env python3
"""Generate Word doc for OpenClaw Day Guide."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.expanduser("~/Desktop/OpenClaw_Day_体验指南.docx")
doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.8)
    s.right_margin = Cm(2.8)

# Colors
NAVY = RGBColor(0x0F, 0x17, 0x2A)
INDIGO = RGBColor(0x4F, 0x46, 0xE5)
DARK = RGBColor(0x1F, 0x2A, 0x37)
GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ORANGE = RGBColor(0xEA, 0x58, 0x0C)
TEAL = RGBColor(0x0D, 0x94, 0x88)

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.font.color.rgb = DARK
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35


def shade(cell, color_hex):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:fill'), color_hex)
    sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)


def run_styled(p, text, size=10.5, bold=False, color=None, italic=False):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = '微软雅黑'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if color:
        r.font.color.rgb = color
    return r


def heading(text, level=1):
    p = doc.add_paragraph()
    sizes = {1: 24, 2: 16, 3: 13}
    colors = {1: NAVY, 2: INDIGO, 3: DARK}
    space_b = {1: 24, 2: 18, 3: 12}
    p.paragraph_format.space_before = Pt(space_b.get(level, 12))
    p.paragraph_format.space_after = Pt(10 if level <= 2 else 6)
    run_styled(p, text, size=sizes.get(level, 12), bold=True, color=colors.get(level, DARK))
    if level == 1:
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '12')
        bot.set(qn('w:color'), '4F46E5')
        bot.set(qn('w:space'), '6')
        pBdr.append(bot)
        pPr.append(pBdr)
    return p


def body(text, bold=False, color=None, italic=False):
    p = doc.add_paragraph()
    run_styled(p, text, bold=bold, color=color, italic=italic)
    return p


def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:color'), '4F46E5')
    left.set(qn('w:space'), '12')
    pBdr.append(left)
    pPr.append(pBdr)
    run_styled(p, text, size=11, bold=True, color=INDIGO)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        run_styled(p, '•  ')
        run_styled(p, bold_prefix, bold=True)
        run_styled(p, text)
    else:
        run_styled(p, '•  ' + text)
    return p


def numbered(num, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after = Pt(4)
    run_styled(p, f'{num}.  ', bold=True, color=INDIGO)
    if bold_prefix:
        run_styled(p, bold_prefix, bold=True)
        run_styled(p, text)
    else:
        run_styled(p, text)
    return p


def code_block(lines):
    """Render a code/prompt block with shaded background."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        # Add shading via paragraph
        pPr = p._element.get_or_add_pPr()
        pShd = OxmlElement('w:shd')
        pShd.set(qn('w:fill'), 'F3F4F6')
        pShd.set(qn('w:val'), 'clear')
        pPr.append(pShd)
        r = p.add_run(line)
        r.font.size = Pt(9.5)
        r.font.name = 'Consolas'
        r.font.color.rgb = DARK


def scenario_box(title, color_hex, items):
    """Create a scenario section with colored left border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:color'), color_hex)
    left.set(qn('w:space'), '10')
    pBdr.append(left)
    pPr.append(pBdr)
    run_styled(p, title, size=13, bold=True, color=RGBColor.from_string(color_hex))


# ════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_styled(p, 'OpenClaw Day', size=40, bold=True, color=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_styled(p, '体 验 指 南', size=30, bold=True, color=INDIGO)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in ['三大场景 · 即学即用', 'BNBU 公用机体验区专用']:
    run_styled(p, line + '\n', size=13, color=GRAY)

doc.add_page_break()

# ════════════════════════════════════════════
# 开始使用
# ════════════════════════════════════════════

heading('开始使用', 1)

numbered(1, '启动 Ubuntu WSL', '打开 Ubuntu — ')
numbered(2, '点击桌面的 openclaw 快捷入口（等效于访问 http://127.0.0.1:18789）', '浏览器登录 — ')
numbered(3, '学院已提供账号，使用手机号登录：13683398795', 'KimiClaw 账号 — ')

doc.add_paragraph()
quote('完成使用后，记得删除临时文件并把《公用机安全须知.md》放回桌面中央。')

body('OpenClaw 生成的成品文件经常会放在 .openclaw 文件夹下，没有计算机基础的同学可能找起来比较不方便，记得要生成某种文件的时候和 OpenClaw 强调成品要放在（Windows 的，而非 WSL 的）桌面上。', italic=True, color=GRAY)

# ════════════════════════════════════════════
# 场景 1
# ════════════════════════════════════════════

heading('场景 1：根据官网生成活动墙', 1)
body('输出格式：HTML', bold=True, color=TEAL)

heading('你需要提供', 3)
bullet('官网 URL（例如 https://www.bnbu.edu.cn/index.html）')
bullet('想突出的条目数量 / 主题（可选）')

heading('你应该对 OpenClaw 说', 3)
code_block([
    '欢迎来到 BNBU 公用机体验区，请帮我完成"官网活动墙"演示。',
    '任务：',
    '1. 用 summarize 抓取 {https://www.bnbu.edu.cn/index.html}',
    '   的公开资讯，挑选 6 条最具代表性的活动或新闻。',
    '2. 先输出一个 Markdown 草稿（含日期、标签、描述、链接）。',
    '3. 然后把这些信息做成精美的 html，',
    '   并在页脚写明数据来源与更新时间。',
    '4. 完成后把成品都保存到 windows 的桌面，',
    '   记住你是在 wsl 环境里运行的。',
])

heading('会用到的 Skills', 3)
bullet('summarize')

# ════════════════════════════════════════════
# 场景 2
# ════════════════════════════════════════════

heading('场景 2：行业速览报告', 1)
body('输出格式：HTML', bold=True, color=TEAL)

heading('你需要提供', 3)
bullet('行业名称（例如"国内量化投资行业"）')
bullet('关注点：玩家 / 监管 / 机会等（可选）')

heading('你应该对 OpenClaw 说', 3)
code_block([
    '欢迎来到 BNBU 公用机体验区，请帮我完成"行业速览报告"演示。',
    '任务：',
    '1. 使用 market-research-agent，围绕 {国内量化投资行业}',
    '   生成一份 3 页以内的 Markdown 报告，',
    '   包含市场规模、主要玩家、机会与风险、下一步建议。',
    '2. 如需引用外部资料，请用 summarize 抓取公开报道并标注出处。',
    '3. 报告完成后导出为精美的 html，并声明所有内容基于公开信息。',
    '4. 将 Markdown 与 html 都存在 windows 的桌面，',
    '   记住你是在 wsl 环境里运行的。',
])

heading('会用到的 Skills', 3)
bullet('market-research-agent')
bullet('summarize')

# ════════════════════════════════════════════
# 场景 3
# ════════════════════════════════════════════

heading('场景 3：学习资料速记包', 1)
body('输出格式：HTML', bold=True, color=TEAL)

heading('你需要提供', 3)
bullet('本地公开资料的文件夹路径')

heading('你应该对 OpenClaw 说', 3)
code_block([
    '欢迎来到 BNBU 公用机体验区，请帮我完成"学习资料速记包"演示。',
    '任务：',
    '1. 使用 summarize 提炼 {资料文件夹路径} 中的内容，',
    '   生成《速记包》Markdown',
    '   （10 个知识点、示例、练习题、5 道测验及答案）。',
    '2. 将 Markdown 导出为精美的 html。',
    '3. 确认全部内容来自公开资料，并在页脚写明来源。',
    '4. 把最终的 Markdown / html 都存在 windows 的桌面，',
    '   记住你是在 wsl 环境里运行的。',
])

heading('会用到的 Skills', 3)
bullet('summarize')
bullet('ppt-generator / mxe')

# ════════════════════════════════════════════
# 开动脑筋
# ════════════════════════════════════════════

heading('开动脑筋：来自己动动手吧！', 1)

body('现在这个 OpenClaw 生成 PDF 的效果不好，所以交付的都是 HTML。你可以把它训练成可以生成精美 PDF 的龙虾吗？')

doc.add_paragraph()
quote('提示：尝试用不同的 Prompt 策略引导 OpenClaw 输出高质量 PDF，这本身就是一次 AI Agent 实战练习！')

# ════════════════════════════════════════════
# Save
# ════════════════════════════════════════════

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
